"""
seed_gold_standard.py

Parses NEXICODE_JavaScript_Question_Bank.docx and loads it into the
database as SyllabusTopic + Question + GoldAnswer rows.

Expected doc structure (confirmed against your actual file):
  For each of the 10 questions:
    "Question N   [EASY|MEDIUM|HARD]"
    "Topic: <topic title>"
    "Syllabus Reference: <text>"
    "Learning Outcome: <text>"
    Heading "Question"      -> question text paragraph(s) follow
    Heading "Marking Scheme and Model Answers"
    A table: header row + 10 rows (Mark, Description, Model Answer code)

Run from inside the 'backend' folder, same as test_claude.py:
    python seed_gold_standard.py

Place NEXICODE_JavaScript_Question_Bank.docx directly in the 'backend'
folder before running (or edit DOCX_PATH below to point at it).

Requires: pip install python-docx
"""

import os
import re
import sys

from app import create_app
from extensions import db
from models import User, Course, SyllabusTopic, Question, GoldAnswer

try:
    import docx
except ImportError:
    print("ERROR: python-docx is not installed.")
    print("Fix: pip install python-docx")
    sys.exit(1)


DOCX_PATH = os.path.join(os.path.dirname(__file__), "NEXICODE_JavaScript_Question_Bank.docx")

# A single Course/tutor to hang all 10 gold-standard SyllabusTopics off of.
# Re-running this script will find and reuse these instead of duplicating them.
GOLD_TUTOR_EMAIL = "gold-standard@nexicode.local"
GOLD_COURSE_TITLE = "NEXICODE Gold Standard Question Bank"
GOLD_COURSE_MODULE_CODE = "G404-GOLD"

QUESTION_HEADER_RE = re.compile(r"^Question\s+(\d+)\s*\[(\w+)\]\s*$")
TOPIC_RE = re.compile(r"^Topic:\s*(.+)$")
SYLLABUS_REF_RE = re.compile(r"^Syllabus Reference:\s*(.+)$")
LEARNING_OUTCOME_RE = re.compile(r"^Learning Outcome:\s*(.+)$")
MARK_CELL_RE = re.compile(r"^(\d+)\s*/\s*10\s*$")


def extract_marking_scheme_key(paragraphs):
    """Pulls the generic 1-3 / 4-6 / 7-9 / 10 rubric text used for every topic."""
    patterns = {
        "1-3": re.compile(r"^1–3 marks:\s*(.+)$"),
        "4-6": re.compile(r"^4–6 marks:\s*(.+)$"),
        "7-9": re.compile(r"^7–9 marks:\s*(.+)$"),
        "10": re.compile(r"^10 marks:\s*(.+)$"),
    }
    found = {}
    for p in paragraphs:
        text = p.text.strip()
        for key, pattern in patterns.items():
            m = pattern.match(text)
            if m:
                found[key] = m.group(1)

    return (
        f"1-3 marks: {found.get('1-3', '')}\n"
        f"4-6 marks: {found.get('4-6', '')}\n"
        f"7-9 marks: {found.get('7-9', '')}\n"
        f"10 marks: {found.get('10', '')}"
    )


def parse_questions(doc):
    """
    Walks the flat paragraph list and returns a list of dicts:
    {number, difficulty, topic, syllabus_reference, learning_outcome, question_text}
    in document order (matching doc.tables order 1:1).
    """
    paragraphs = doc.paragraphs
    questions = []
    current = None
    state = None  # None | "awaiting_question_text" | "in_question_text"

    for p in paragraphs:
        text = p.text.strip()
        style = p.style.name if p.style else ""

        if not text:
            continue

        header_match = QUESTION_HEADER_RE.match(text)
        if header_match:
            if current:
                questions.append(current)
            current = {
                "number": int(header_match.group(1)),
                "difficulty": header_match.group(2).lower(),
                "topic": None,
                "syllabus_reference": None,
                "learning_outcome": None,
                "question_text": "",
            }
            state = None
            continue

        if current is None:
            continue  # still in the document's front-matter, before Question 1

        topic_match = TOPIC_RE.match(text)
        if topic_match:
            current["topic"] = topic_match.group(1)
            continue

        ref_match = SYLLABUS_REF_RE.match(text)
        if ref_match:
            current["syllabus_reference"] = ref_match.group(1)
            continue

        outcome_match = LEARNING_OUTCOME_RE.match(text)
        if outcome_match:
            current["learning_outcome"] = outcome_match.group(1)
            continue

        if style == "Heading 2" and text == "Question":
            state = "in_question_text"
            continue

        if style == "Heading 2" and text == "Marking Scheme and Model Answers":
            state = None
            continue

        if state == "in_question_text":
            current["question_text"] = (current["question_text"] + " " + text).strip()
            continue

    if current:
        questions.append(current)

    return questions


def parse_gold_answers(table):
    """Returns a list of (gold_mark, answer_text) tuples from an 11-row table."""
    answers = []
    for row in table.rows[1:]:  # skip header row
        mark_cell = row.cells[0].text.strip()
        answer_cell = row.cells[2].text.strip()
        mark_match = MARK_CELL_RE.match(mark_cell)
        if not mark_match:
            print(f"  WARNING: couldn't parse mark from cell '{mark_cell}', skipping row.")
            continue
        answers.append((int(mark_match.group(1)), answer_cell))
    return answers


def get_or_create_gold_tutor():
    tutor = User.query.filter_by(email=GOLD_TUTOR_EMAIL).first()
    if tutor:
        return tutor

    from werkzeug.security import generate_password_hash
    import uuid as uuid_lib

    print(f"No gold-standard tutor user found — creating one ({GOLD_TUTOR_EMAIL}).")
    tutor = User(
        name="NEXICODE Gold Standard Seed",
        email=GOLD_TUTOR_EMAIL,
        password_hash=generate_password_hash(uuid_lib.uuid4().hex),  # random, unused password
        role="tutor",
    )
    db.session.add(tutor)
    db.session.commit()
    return tutor


def get_or_create_gold_course(tutor):
    course = Course.query.filter_by(module_code=GOLD_COURSE_MODULE_CODE).first()
    if course:
        return course

    print(f"Creating course '{GOLD_COURSE_TITLE}'.")
    course = Course(
        title=GOLD_COURSE_TITLE,
        module_code=GOLD_COURSE_MODULE_CODE,
        tutor_id=tutor.id,
    )
    db.session.add(course)
    db.session.commit()
    return course


def get_or_create_topic(course, topic_title, learning_outcome, marking_rubric):
    topic = SyllabusTopic.query.filter_by(
        course_id=course.id, topic_title=topic_title
    ).first()
    if topic:
        return topic

    topic = SyllabusTopic(
        course_id=course.id,
        topic_title=topic_title,
        learning_outcomes=learning_outcome or "",
        marking_rubric=marking_rubric,
    )
    db.session.add(topic)
    db.session.commit()
    return topic


def get_or_create_question(topic, question_text, difficulty):
    question = Question.query.filter_by(
        topic_id=topic.id, question_text=question_text
    ).first()
    if question:
        return question, False

    question = Question(
        topic_id=topic.id,
        question_text=question_text,
        difficulty=difficulty,
    )
    db.session.add(question)
    db.session.commit()
    return question, True


def main():
    if not os.path.exists(DOCX_PATH):
        print(f"ERROR: could not find {DOCX_PATH}")
        print("Place NEXICODE_JavaScript_Question_Bank.docx in the backend folder, or edit DOCX_PATH.")
        sys.exit(1)

    doc = docx.Document(DOCX_PATH)
    marking_rubric = extract_marking_scheme_key(doc.paragraphs)
    parsed_questions = parse_questions(doc)
    tables = doc.tables

    if len(parsed_questions) != len(tables):
        print(f"WARNING: found {len(parsed_questions)} question headers but {len(tables)} tables.")
        print("These are expected to match 1:1 in document order. Double check the .docx before trusting the results.")

    app = create_app()
    with app.app_context():
        tutor = get_or_create_gold_tutor()
        course = get_or_create_gold_course(tutor)

        questions_added = 0
        questions_reused = 0
        gold_answers_added = 0
        gold_answers_skipped = 0

        for q_data, table in zip(parsed_questions, tables):
            print(f"Question {q_data['number']} [{q_data['difficulty']}] — {q_data['topic']}")

            topic = get_or_create_topic(
                course,
                topic_title=q_data["topic"] or f"Untitled Topic {q_data['number']}",
                learning_outcome=q_data["learning_outcome"],
                marking_rubric=marking_rubric,
            )

            question, was_created = get_or_create_question(
                topic, q_data["question_text"], q_data["difficulty"]
            )
            if was_created:
                questions_added += 1
            else:
                questions_reused += 1

            gold_answers = parse_gold_answers(table)
            for gold_mark, answer_text in gold_answers:
                existing = GoldAnswer.query.filter_by(
                    question_id=question.id, gold_mark=gold_mark
                ).first()
                if existing:
                    gold_answers_skipped += 1
                    continue

                gold_answer = GoldAnswer(
                    question_id=question.id,
                    answer_text=answer_text,
                    gold_mark=gold_mark,
                )
                db.session.add(gold_answer)
                gold_answers_added += 1

            db.session.commit()

        print()
        print(f"Questions added: {questions_added} (reused {questions_reused} already existing)")
        print(f"Gold answers added: {gold_answers_added} (skipped {gold_answers_skipped} already existing)")


if __name__ == "__main__":
    main()
